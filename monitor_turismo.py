import os
import feedparser
from openai import OpenAI
from datetime import datetime, timedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import quote

# ══════════════════════════════════════════════════════
# FUENTES RSS OFICIALES Y ESPECIALIZADAS
# ══════════════════════════════════════════════════════
FUENTES_OFICIALES = [
    ("UN Tourism (UNWTO)",      "https://www.unwto.org/feed"),
    ("WTTC",                    "https://wttc.org/feed"),
    ("IATA News",               "https://www.iata.org/en/pressroom/pr/rss/"),
    ("OECD Tourism",            "https://www.oecd-ilibrary.org/rss/content/topic/tourism"),
    ("PATA",                    "https://www.pata.org/feed/"),
]

FUENTES_ESPECIALIZADAS = [
    ("Skift",                   "https://skift.com/feed/"),
    ("PhocusWire",              "https://www.phocuswire.com/rss"),
    ("Hosteltur",               "https://www.hosteltur.com/rss"),
    ("Travel Weekly",           "https://www.travelweekly.com/rss"),
    ("TTG Media",               "https://www.ttgmedia.com/rss"),
    ("Travel Daily Media",      "https://www.traveldailymedia.com/feed/"),
    ("Hotel News Now",          "https://www.hotelnewsnow.com/rss"),
    ("TTN Worldwide",           "https://www.ttnworldwide.com/rss.xml"),
]

# ══════════════════════════════════════════════════════
# GOOGLE NEWS RSS — PALABRAS CLAVE (prensa general mundial)
# ══════════════════════════════════════════════════════
KEYWORDS_EN = [
    "international tourism",
    "travel industry",
    "tourism policy",
    "airline passengers",
    "hotel occupancy",
    "sustainable tourism",
    "cruise industry",
    "visa travel",
]

KEYWORDS_ES = [
    "turismo mundial",
    "industria del turismo",
    "turismo latinoamérica",
    "turismo Chile",
    "viajeros internacionales",
]

def google_news_url(keyword, lang="en", country="US"):
    q = quote(keyword)
    return f"https://news.google.com/rss/search?q={q}&hl={lang}&gl={country}&ceid={country}:{lang}"

FUENTES_GOOGLE_NEWS = (
    [(kw, google_news_url(kw, "en", "US")) for kw in KEYWORDS_EN] +
    [(kw, google_news_url(kw, "es", "CL")) for kw in KEYWORDS_ES]
)

# ══════════════════════════════════════════════════════
# RECOLECCIÓN Y DEDUPLICACIÓN
# ══════════════════════════════════════════════════════
def recolectar_noticias():
    todas = []
    titulos_vistos = set()
    hace_7_dias = datetime.now() - timedelta(days=7)

    categorias = [
        ("🏛️ Organismos Oficiales",  FUENTES_OFICIALES),
        ("📰 Medios Especializados", FUENTES_ESPECIALIZADAS),
        ("🌐 Prensa General",        FUENTES_GOOGLE_NEWS),
    ]

    for categoria, fuentes in categorias:
        for nombre, url in fuentes:
            try:
                feed = feedparser.parse(url)
                count = 0
                for entry in feed.entries[:8]:
                    titulo = entry.get("title", "").strip()
                    if not titulo or titulo.lower() in titulos_vistos:
                        continue
                    titulos_vistos.add(titulo.lower())
                    todas.append({
                        "titulo":    titulo,
                        "resumen":   entry.get("summary", "")[:400],
                        "fuente":    nombre,
                        "link":      entry.get("link", ""),
                        "categoria": categoria,
                    })
                    count += 1
                    if count >= 5:
                        break
                print(f"  ✓ {nombre}: {count} noticias")
            except Exception as e:
                print(f"  ✗ {nombre}: {e}")

    print(f"\n→ Total noticias únicas recolectadas: {len(todas)}\n")
    return todas

# ══════════════════════════════════════════════════════
# GENERACIÓN DEL RESUMEN CON GPT
# ══════════════════════════════════════════════════════
def generar_resumen(noticias):
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    noticias_texto = "\n\n".join([
        f"[{n['categoria']} | {n['fuente']}]\n{n['titulo']}\n{n['resumen']}"
        for n in noticias
    ])

    prompt = f"""Eres un analista senior de turismo internacional que prepara informes para la Subsecretaria de Turismo de Chile.

Tienes acceso a {len(noticias)} noticias de turismo de la última semana, provenientes de organismos oficiales, medios especializados y prensa general mundial.

NOTICIAS:
{noticias_texto}

Elabora un RESUMEN EJECUTIVO SEMANAL en español con esta estructura exacta:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. TITULARES DE LA SEMANA
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Selecciona las 5 noticias más relevantes globalmente. Para cada una:
- Título claro y directo
- Párrafo de 2-3 oraciones explicando qué pasó y por qué importa
- Fuente entre paréntesis

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2. TENDENCIAS EMERGENTES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Identifica 2-3 patrones o temas recurrentes en las noticias de esta semana.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
3. RELEVANCIA PARA CHILE Y LATINOAMÉRICA
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Analiza qué implican estos eventos para el turismo receptor de Chile, 
la competitividad regional y posibles oportunidades o amenazas.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
4. PARA TENER EN RADAR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2 temas o eventos a monitorear la próxima semana, con breve justificación.

Tono: profesional, directo, orientado a decisiones de política pública. Máximo 900 palabras."""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        max_tokens=2000,
        messages=[
            {"role": "system", "content": "Eres un analista experto en turismo internacional. Escribes en español formal chileno."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content

# ══════════════════════════════════════════════════════
# GENERACIÓN DEL ARCHIVO WORD
# ══════════════════════════════════════════════════════
def crear_word(resumen, noticias):
    doc = Document()

    # Márgenes
    from docx.shared import Inches, Pt, Cm
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

    # Encabezado institucional
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("SUBSECRETARÍA DE TURISMO DE CHILE")
    run.bold = True
    run.font.size = Pt(11)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Unidad de Estudios").bold = True

    doc.add_paragraph("")

    # Título principal
    titulo = doc.add_heading("MONITOR DE TURISMO MUNDIAL", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha = datetime.now().strftime("%d de %B de %Y")
    semana = datetime.now().strftime("Semana N° %W — %Y")
    subtitulo = doc.add_heading(f"Resumen Ejecutivo Semanal\n{semana}", 2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run("Fecha de elaboración: ").bold = True
    meta.add_run(fecha)

    meta2 = doc.add_paragraph()
    meta2.add_run("Noticias analizadas: ").bold = True
    meta2.add_run(f"{len(noticias)} artículos de {len(set(n['fuente'] for n in noticias))} fuentes")

    meta3 = doc.add_paragraph()
    meta3.add_run("Fuentes: ").bold = True
    meta3.add_run("UN Tourism, WTTC, IATA, OECD, Skift, PhocusWire, Hosteltur, Travel Weekly, TTG Media, Google News (EN/ES) y otras")

    doc.add_paragraph("─" * 65)
    doc.add_paragraph("")

    # Contenido del resumen
    for linea in resumen.split("\n"):
        linea = linea.strip()
        if not linea:
            doc.add_paragraph("")
            continue
        if linea.startswith("━"):
            continue
        if linea and linea[0].isdigit() and "." in linea[:3]:
            doc.add_heading(linea, level=2)
        elif linea.startswith("-") or linea.startswith("•"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(linea.lstrip("-• "))
        else:
            doc.add_paragraph(linea)

    # Pie
    doc.add_paragraph("")
    doc.add_paragraph("─" * 65)
    pie = doc.add_paragraph()
    pie.add_run(
        f"Documento generado automáticamente mediante IA (GPT-4o mini + RSS feeds). "
        f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')} hrs."
    ).italic = True

    nombre = f"Monitor_Turismo_{datetime.now().strftime('%Y_%m_%d')}.docx"
    doc.save(nombre)
    print(f"✅ Archivo Word generado: {nombre}")
    return nombre

# ══════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 55)
    print("  MONITOR DE TURISMO MUNDIAL")
    print(f"  {datetime.now().strftime('%d/%m/%Y %H:%M')} hrs")
    print("=" * 55)

    print("\n🔍 Recolectando noticias...")
    noticias = recolectar_noticias()

    print("🤖 Generando resumen ejecutivo con GPT-4o mini...")
    resumen = generar_resumen(noticias)

    print("📄 Creando archivo Word...")
    crear_word(resumen, noticias)

    print("\n🎉 ¡Listo! Monitor generado exitosamente.")
    
